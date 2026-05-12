"""
KazLLM Annotator — локальное веб-приложение для аннотаторов
Института языкознания им. А. Байтурсынова.

Назначение: создавать единицы данных по инструкции v3 (2026-04-21)
для дообучения Qwen3.5 на казахском языке.

Запуск:
    pip install -r requirements.txt
    python app.py
    Открыть http://127.0.0.1:5000

Каждый аннотатор ставит приложение себе на компьютер.
Данные хранятся в локальном файле data.db (SQLite).
"""

import os
import re
import json
import hashlib
import secrets
import sqlite3
import datetime as dt
from io import BytesIO
from functools import wraps

from flask import (
    Flask, request, session, redirect, url_for, render_template,
    flash, jsonify, send_file, g, abort
)
from werkzeug.security import generate_password_hash, check_password_hash
from docx import Document
from docx.oxml.ns import qn

from i18n import t as _t, TRANSLATIONS, LANGS, DEFAULT_LANG
from tools_schema import (
    TOOL_SCHEMAS, empty_input_for, empty_output_for, check_value_type,
    schema_for_form, get_input_spec, get_output_keys,
)
from examples_data import EXAMPLES, get_example, list_examples

# ─────────────────────────────────────────────────────────────────────────────
# Конфигурация
# ─────────────────────────────────────────────────────────────────────────────

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DB_PATH = os.path.join(BASE_DIR, "data.db")

app = Flask(__name__)
app.config["SECRET_KEY"] = os.environ.get("SECRET_KEY") or secrets.token_hex(32)
app.config["MAX_CONTENT_LENGTH"] = 5 * 1024 * 1024  # 5 МБ

# ─────────────────────────────────────────────────────────────────────────────
# Справочники из инструкции v3
# ─────────────────────────────────────────────────────────────────────────────

ARCHETYPES = [
    ("dialogue",          "Диалог (обычный многотуровый)"),
    ("preference",        "Пара предпочтений (DPO)"),
    ("cultural_reasoning","Культурное рассуждение"),
    ("business_scenario", "Бизнес-сценарий с инструментами"),
    ("benchmark",         "Бенчмарк (удерживаемый)"),
]

DOMAINS = [
    "banking", "egov", "tax", "fsms", "enbek", "logistics",
    "education", "health", "commerce", "legal", "culture", "other",
]

CULTURAL_ANCHORS = [
    "national_value", "philosophy", "tradition",
    "business_norm", "regional", "modern_life", "religion",
]

COMPLEXITIES = ["simple", "moderate", "complex"]

LICENSES = [
    "public-domain", "cc-by", "cc-by-sa", "cc-by-nc",
    "original-institute", "written-permission",
]

REJECTION_REASONS = [
    ("morphological-error",   "Морфологическая ошибка (падеж, согласование)"),
    ("factual-error",         "Фактическая ошибка"),
    ("cultural-error",        "Культурная ошибка"),
    ("tool-error",            "Ошибка вызова инструмента"),
    ("hallucinated-citation", "Галлюцинация ссылки"),
    ("refusal-mismatch",      "Несоответствие отказа"),
    ("reasoning-error",       "Ошибка рассуждения"),
    ("tone-mismatch",         "Несоответствие тона / регистра"),
]

# Библиотека инструментов (Приложение А, v1) — 55 инструментов
TOOL_LIBRARY = [
    # Банкинг
    "kaspi.balance", "kaspi.transfer", "kaspi.pay", "kaspi.dispute",
    "halyk.balance", "halyk.transfer", "halyk.card_block",
    "forte.balance", "jusan.balance",
    # egov.kz
    "egov.login", "egov.service_request", "egov.status",
    "egov.gbdfl_lookup", "egov.verify_iin", "egov.get_certificate",
    "egov.law_lookup",
    # КГД / Налоги
    "kgd.tax_status", "kgd.declaration_status", "kgd.deadline", "kgd.bin_lookup",
    # ФСМС
    "fsms.status", "fsms.claim_submit",
    # Труд
    "enbek.search_job", "enbek.apply", "enbek.employer_stats",
    # Почта / логистика
    "kazpost.track", "kazpost.estimate",
    # Образование
    "ent.register", "edu.university_faq", "edu.school_schedule",
    # Здравоохранение
    "health.find_clinic", "health.book_appointment", "health.prescription_status",
    # Коммерция
    "shop.search", "shop.order_status", "shop.return_request",
    # События
    "events.search", "events.book_ticket",
    # Путешествия
    "flights.search", "trains.search", "hotel.search",
    # Связь
    "comm.sms", "comm.email", "comm.call",
    # Карты
    "maps.geocode", "maps.directions", "maps.nearby",
    # Поиск
    "search.web", "search.news", "search.wiki", "search.docs",
    # Расписание
    "calendar.create_event", "calendar.availability",
    # Система
    "system.now", "system.translate",
]
TOOL_SET = set(TOOL_LIBRARY)

# ─────────────────────────────────────────────────────────────────────────────
# База данных (SQLite)
# ─────────────────────────────────────────────────────────────────────────────

def get_db():
    if "db" not in g:
        g.db = sqlite3.connect(DB_PATH)
        g.db.row_factory = sqlite3.Row
        g.db.execute("PRAGMA foreign_keys = ON")
    return g.db


@app.teardown_appcontext
def close_db(_exc):
    db = g.pop("db", None)
    if db is not None:
        db.close()


def init_db():
    """Создаёт таблицы при первом запуске."""
    conn = sqlite3.connect(DB_PATH)
    conn.executescript("""
    CREATE TABLE IF NOT EXISTS users (
        id            INTEGER PRIMARY KEY AUTOINCREMENT,
        username      TEXT UNIQUE NOT NULL,
        full_name     TEXT NOT NULL,
        password_hash TEXT NOT NULL,
        created_at    TEXT NOT NULL DEFAULT (datetime('now'))
    );

    CREATE TABLE IF NOT EXISTS units (
        id              INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id         INTEGER NOT NULL REFERENCES users(id) ON DELETE CASCADE,
        unit_code       TEXT NOT NULL,
        archetype       TEXT NOT NULL,
        domain          TEXT,
        cultural_anchor TEXT,
        business_use_case TEXT,
        complexity      TEXT,
        tools_in_scope  TEXT,
        scenario_kk     TEXT,
        turns_json      TEXT NOT NULL,        -- JSON-список туров
        preference_json TEXT,                  -- JSON для preference-архетипа
        source          TEXT NOT NULL,
        license         TEXT NOT NULL,
        reviewer        TEXT,
        started_at      TEXT NOT NULL,
        completed_at    TEXT NOT NULL,
        created_at      TEXT NOT NULL DEFAULT (datetime('now'))
    );

    CREATE INDEX IF NOT EXISTS idx_units_user ON units(user_id);
    """)
    conn.commit()
    conn.close()


# Инициализируем БД при импорте — это нужно и для тестов,
# и для продакшен-WSGI запуска без if __name__ == '__main__'.
init_db()


# ─────────────────────────────────────────────────────────────────────────────
# Авторизация
# ─────────────────────────────────────────────────────────────────────────────

def login_required(view):
    @wraps(view)
    def wrapped(*args, **kwargs):
        if "user_id" not in session:
            return redirect(url_for("login"))
        return view(*args, **kwargs)
    return wrapped


def current_user():
    if "user_id" not in session:
        return None
    db = get_db()
    return db.execute(
        "SELECT * FROM users WHERE id = ?", (session["user_id"],)
    ).fetchone()


@app.context_processor
def inject_user():
    return {"user": current_user()}


# ─────────────────────────────────────────────────────────────────────────────
# i18n
# ─────────────────────────────────────────────────────────────────────────────

def get_lang() -> str:
    """Берёт язык из cookie. По умолчанию — казахский."""
    lang = request.cookies.get("lang", DEFAULT_LANG)
    return lang if lang in LANGS else DEFAULT_LANG


@app.context_processor
def inject_i18n():
    lang = get_lang()
    return {
        "lang": lang,
        "other_lang": "ru" if lang == "kk" else "kk",
        "t": lambda key, *args: _t(key, lang, *args),
    }


@app.route("/lang/<code>")
def set_lang(code):
    if code not in LANGS:
        code = DEFAULT_LANG
    nxt = request.args.get("next") or url_for("index")
    # Безопасность: разрешаем только относительные URL внутри сайта
    if not nxt.startswith("/"):
        nxt = url_for("index")
    resp = redirect(nxt)
    resp.set_cookie("lang", code, max_age=60 * 60 * 24 * 365, samesite="Lax")
    return resp


def tt(key: str, *args) -> str:
    """Сокращение для использования в Python-коде (валидация и т.п.)."""
    return _t(key, get_lang(), *args)


def _localized_archetypes():
    return [(v, tt(f"archetype.{v}")) for v, _ in ARCHETYPES]


def _localized_reasons():
    return [(v, tt(f"reason.{v}")) for v, _ in REJECTION_REASONS]


# ─────────────────────────────────────────────────────────────────────────────
# Валидация (защита «от тупого»)
# ─────────────────────────────────────────────────────────────────────────────

USERNAME_RE = re.compile(r"^[a-zA-Z0-9_.\-]{3,32}$")
UNIT_CODE_RE = re.compile(r"^ly-ds-2026h2-[a-z0-9_]+-\d{3,6}$")
ISO_DATETIME_RE = re.compile(
    r"^\d{4}-\d{2}-\d{2}T\d{2}:\d{2}:\d{2}([+-]\d{2}:?\d{2}|Z)?$"
)
IIN_RE = re.compile(r"^\d{12}$")
DATE_ISO_RE = re.compile(r"^\d{4}-\d{2}-\d{2}$")


def validate_iin(iin: str) -> bool:
    """Проверяет ИИН Казахстана: 12 цифр + контрольная сумма по ГОСТ.
    Используется в проверке аргументов вызовов инструментов.

    Тримит пробелы, защищая от случайных whitespace-ов при копировании.
    """
    if not isinstance(iin, str):
        return False
    iin = iin.strip()
    if not IIN_RE.match(iin):
        return False
    digits = [int(c) for c in iin]
    w1 = [1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11]
    w2 = [3, 4, 5, 6, 7, 8, 9, 10, 11, 1, 2]
    s = sum(d * w for d, w in zip(digits[:11], w1))
    check = s % 11
    if check == 10:
        s = sum(d * w for d, w in zip(digits[:11], w2))
        check = s % 11
        if check == 10:
            return False
    return check == digits[11]


def diagnose_iin(iin: str) -> str | None:
    """Возвращает None если ИИН валиден, иначе строку-диагноз: что именно не так.
    Используется для более понятных сообщений об ошибках валидации."""
    if not isinstance(iin, str):
        return f"ИИН должен быть строкой, получен {type(iin).__name__}"
    original = iin
    iin = iin.strip()
    if iin != original:
        # Пробелы — но мы их триммим в validate_iin, поэтому это уже не блокер.
        # Просто продолжаем проверку дальше с обрезанной строкой.
        pass
    if len(iin) == 0:
        return "ИИН пустой"
    if len(iin) != 12:
        return f"ИИН должен быть ровно 12 цифр, получено {len(iin)} символов: {iin!r}"
    # Проверка на нецифровые символы — диагностика частых случаев
    non_digits = [(i, c) for i, c in enumerate(iin) if not c.isdigit()]
    if non_digits:
        # Пытаемся понять, кириллица ли это (О/о, З/з, И/и)
        cyrillic_lookalikes = {"О": "0", "о": "0", "З": "3", "з": "3",
                                "И": "1", "и": "1", "Н": "Н?", "В": "В?"}
        hints = []
        for i, c in non_digits:
            if c in cyrillic_lookalikes:
                hints.append(f"позиция {i+1}: «{c}» (похоже на кириллицу — нужна цифра {cyrillic_lookalikes[c]})")
            else:
                hints.append(f"позиция {i+1}: «{c}»")
        return "ИИН содержит нецифровые символы — " + "; ".join(hints)
    # Здесь все 12 цифр, значит проблема только в контрольной сумме
    digits = [int(c) for c in iin]
    w1 = [1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11]
    w2 = [3, 4, 5, 6, 7, 8, 9, 10, 11, 1, 2]
    s = sum(d * w for d, w in zip(digits[:11], w1))
    check = s % 11
    if check == 10:
        s = sum(d * w for d, w in zip(digits[:11], w2))
        check = s % 11
        if check == 10:
            return f"ИИН {iin!r}: невозможно вычислить контрольную сумму (оба алгоритма дают остаток 10 — такой ИИН не выдаётся)"
    if check != digits[11]:
        return f"ИИН {iin!r}: контрольная сумма не сходится (ожидалась цифра {check} на 12-й позиции, а стоит {digits[11]}). Проверьте на опечатку."
    return None


def parse_iso_dt(s: str):
    """Парсит ISO-8601 timestamp с таймзоной."""
    if not s or not ISO_DATETIME_RE.match(s):
        return None
    try:
        # Python поддерживает 'Z' начиная с 3.11 в fromisoformat,
        # но для совместимости преобразуем
        s2 = s.replace("Z", "+00:00")
        return dt.datetime.fromisoformat(s2)
    except ValueError:
        return None


def validate_unit(form_data: dict, turns: list, preference: dict | None) -> list[str]:
    """
    Валидирует единицу по инструкции v3.
    Возвращает список ошибок (пустой = всё ок).
    Сообщения локализованы по текущему cookie-языку.
    """
    errors: list[str] = []

    # ── Базовые поля
    archetype = form_data.get("archetype", "").strip()
    if archetype not in {a[0] for a in ARCHETYPES}:
        errors.append(tt("v.archetype_bad"))

    # Формат unit_code (серверная валидация — браузерный pattern легко обходится)
    unit_code = form_data.get("unit_code", "").strip()
    if not UNIT_CODE_RE.match(unit_code):
        errors.append(tt("v.unit_code_format"))

    domain = form_data.get("domain", "").strip()
    if domain and domain not in DOMAINS:
        errors.append(tt("v.domain_bad", domain))

    complexity = form_data.get("complexity", "").strip()
    if complexity and complexity not in COMPLEXITIES:
        errors.append(tt("v.complexity_bad"))

    cultural_anchor = form_data.get("cultural_anchor", "").strip()
    if cultural_anchor and cultural_anchor not in CULTURAL_ANCHORS:
        errors.append(tt("v.anchor_bad"))

    license_v = form_data.get("license", "").strip()
    if license_v not in LICENSES:
        errors.append(tt("v.license_required"))

    source = form_data.get("source", "").strip()
    if not source:
        errors.append(tt("v.source_required"))
    elif source.lower() == "unknown":
        errors.append(tt("v.source_unknown"))

    # ── Метки времени
    started_raw = form_data.get("started_at", "").strip()
    completed_raw = form_data.get("completed_at", "").strip()
    started = parse_iso_dt(started_raw)
    completed = parse_iso_dt(completed_raw)
    if not started:
        errors.append(tt("v.started_bad"))
    if not completed:
        errors.append(tt("v.completed_bad"))
    if started and completed:
        delta_min = (completed - started).total_seconds() / 60.0
        if delta_min < 0:
            errors.append(tt("v.completed_before"))
        elif delta_min < 3:
            errors.append(tt("v.too_fast", delta_min))
        elif delta_min > 45:
            errors.append(tt("v.too_slow", delta_min))

    # ── Архетипо-зависимые требования
    tools_scope_raw = form_data.get("tools_in_scope", "").strip()

    if archetype == "business_scenario":
        if not form_data.get("business_use_case", "").strip():
            errors.append(tt("v.business_buc"))
        if len(turns) < 3:
            errors.append(tt("v.business_min_turns"))
        if len(turns) > 8:
            errors.append(tt("v.business_max_turns"))
        any_tool = any(t.get("tool_name") for t in turns if t["role"] == "ASSISTANT")
        if not any_tool:
            errors.append(tt("v.business_no_tool"))
        # П.7: TOOLS_IN_SCOPE обязателен для business_scenario
        if not tools_scope_raw:
            errors.append(tt("v.business_scope_required"))

    if archetype == "cultural_reasoning":
        if not cultural_anchor:
            errors.append(tt("v.cultural_anchor"))
        for i, t in enumerate(turns, 1):
            if t["role"] == "ASSISTANT" and not t.get("thinking_kk", "").strip():
                errors.append(tt("v.cultural_thinking", i))

    if archetype == "preference":
        if not preference:
            errors.append(tt("v.pref_block"))
        else:
            try:
                gti = int(preference.get("good_turn_idx", 0))
            except (TypeError, ValueError):
                gti = 0
            if gti < 1 or gti > len(turns):
                errors.append(tt("v.pref_idx_range", gti, len(turns)))
            elif turns[gti - 1]["role"] != "ASSISTANT":
                errors.append(tt("v.pref_idx_role", gti))
            else:
                # П.12: bad_alternative не должна быть равна good content (тривиальная пара)
                good_content = (turns[gti - 1].get("content_kk") or "").strip()
                bad_content = (preference.get("bad_alternative_kk") or "").strip()
                if good_content and bad_content and _normalize_text(good_content) == _normalize_text(bad_content):
                    errors.append(tt("v.pref_identical"))
            if not preference.get("bad_alternative_kk", "").strip():
                errors.append(tt("v.pref_bad_required"))
            br = preference.get("bad_reason", "")
            if br not in {r[0] for r in REJECTION_REASONS}:
                errors.append(tt("v.pref_reason"))

    if archetype == "benchmark":
        if not any(
            (t.get("content_kk", "").strip())
            for t in turns if t["role"] == "ASSISTANT"
        ):
            errors.append(tt("v.bench_no_answer"))

    # ── Туры
    if not turns:
        errors.append(tt("v.no_turns"))
    role_seen_user = False
    used_tools: list[str] = []  # для проверки соответствия со scope
    last_assistant_mock: dict | None = None  # для проверки TOOL_RESULT==mock
    last_assistant_tool_name: str | None = None
    prev_role: str | None = None  # для проверки чередования ролей

    for i, t in enumerate(turns, 1):
        role = t.get("role", "")
        if role not in ("USER", "ASSISTANT", "TOOL_RESULT"):
            errors.append(tt("v.turn_role_bad", i, role))
            prev_role = role
            continue

        # П.6: чередование ролей
        # USER → USER нельзя; ASSISTANT → ASSISTANT нельзя без TOOL_RESULT между ними
        if prev_role == "USER" and role == "USER":
            errors.append(tt("v.role_user_user", i))
        if prev_role == "TOOL_RESULT" and role == "USER":
            # после TOOL_RESULT ожидаем ASSISTANT (он реагирует на результат)
            errors.append(tt("v.role_after_tool_result", i))

        if role == "USER":
            role_seen_user = True
            if not t.get("content_kk", "").strip():
                errors.append(tt("v.turn_user_empty", i))
            last_assistant_mock = None
            last_assistant_tool_name = None

        elif role == "ASSISTANT":
            content_ok = bool(t.get("content_kk", "").strip())
            tool_ok = bool(t.get("tool_name", "").strip())
            if not content_ok and not tool_ok:
                errors.append(tt("v.turn_assistant_empty", i))
            if tool_ok:
                tn = t["tool_name"].strip()
                if tn not in TOOL_SET:
                    errors.append(tt("v.tool_unknown", i, tn))
                    last_assistant_mock = None
                    last_assistant_tool_name = None
                else:
                    used_tools.append(tn)
                    last_assistant_tool_name = tn
                    schema = TOOL_SCHEMAS.get(tn)

                    # П.3: cross-domain warning
                    if schema and domain and domain not in schema.get("domain", []):
                        errors.append(tt("v.tool_domain_mismatch", i, tn, domain,
                                         ", ".join(schema.get("domain", []))))

                    # Валидация args по схеме
                    args_raw = t.get("tool_args", "").strip()
                    parsed_args = None
                    if args_raw:
                        try:
                            parsed_args = json.loads(args_raw)
                        except json.JSONDecodeError as e:
                            errors.append(tt("v.tool_args_json", i, e.msg))

                    if parsed_args is not None:
                        if not isinstance(parsed_args, dict):
                            errors.append(tt("v.tool_args_not_obj", i))
                        elif schema:
                            input_spec = get_input_spec(tn) or {}
                            required = set(input_spec.get("required", []))
                            optional = set(input_spec.get("optional", []))
                            allowed = required | optional
                            types = input_spec.get("types", {})
                            keys = set(parsed_args.keys())

                            missing = required - keys
                            for mk in sorted(missing):
                                errors.append(tt("v.tool_arg_missing", i, tn, mk))

                            unexpected = keys - allowed
                            for uk in sorted(unexpected):
                                errors.append(tt("v.tool_arg_unexpected", i, tn, uk))

                            for k, v in parsed_args.items():
                                if k in types:
                                    err = check_value_type(v, types[k], validate_iin)
                                    if err:
                                        # Если это была проверка ИИН — даём диагностический текст
                                        if types[k] == "iin" and isinstance(v, str):
                                            diag = diagnose_iin(v)
                                            if diag:
                                                errors.append(tt("v.tool_arg_type", i, k, diag))
                                                continue
                                        errors.append(tt("v.tool_arg_type", i, k, err))
                        else:
                            # Схемы нет — fallback: эвристика по ключу (как было раньше)
                            for k, v in parsed_args.items():
                                if "iin" in k.lower() and isinstance(v, str):
                                    if not validate_iin(v):
                                        diag = diagnose_iin(v) or "ИИН некорректен"
                                        errors.append(tt("v.tool_args_iin_diag", i, diag))

                    # mock_result
                    mock_raw = t.get("tool_mock_result", "").strip()
                    if mock_raw:
                        try:
                            parsed_mock = json.loads(mock_raw)
                            last_assistant_mock = parsed_mock if isinstance(parsed_mock, dict) else None
                            # Опциональная проверка: ключи mock соответствуют output-схеме
                            if schema and isinstance(parsed_mock, dict):
                                output_keys = get_output_keys(tn)
                                if output_keys:
                                    extra_keys = set(parsed_mock.keys()) - output_keys
                                    for ek in sorted(extra_keys):
                                        errors.append(tt("v.tool_mock_extra_key", i, tn, ek))
                        except json.JSONDecodeError:
                            errors.append(tt("v.tool_mock_json", i))
                            last_assistant_mock = None
                    else:
                        last_assistant_mock = None
            else:
                last_assistant_mock = None
                last_assistant_tool_name = None

        elif role == "TOOL_RESULT":
            content_raw = t.get("content_kk", "").strip()
            if not content_raw:
                errors.append(tt("v.tool_result_empty", i))
            else:
                # TOOL_RESULT теперь должен быть валидным JSON
                try:
                    parsed_tr = json.loads(content_raw)
                    # П.1: TOOL_RESULT ↔ expected mock_result
                    if last_assistant_mock is not None and isinstance(parsed_tr, dict):
                        if parsed_tr != last_assistant_mock:
                            errors.append(tt("v.tool_result_mock_mismatch", i))
                except json.JSONDecodeError as e:
                    errors.append(tt("v.tool_result_not_json", i, e.msg))

        prev_role = role

    if not role_seen_user:
        errors.append(tt("v.no_user_turn"))

    # TOOL_RESULT правильно следует за ASSISTANT-с-инструментом
    for i, t in enumerate(turns):
        if t.get("role") == "TOOL_RESULT":
            if i == 0 or turns[i - 1].get("role") != "ASSISTANT":
                errors.append(tt("v.tool_result_after", i + 1))
            elif not turns[i - 1].get("tool_name"):
                errors.append(tt("v.tool_result_no_call", i + 1))

    # ── Tools in scope: соответствие фактическому использованию (П.2)
    if tools_scope_raw:
        scope = [s.strip() for s in tools_scope_raw.split(",") if s.strip()]
        for tn in scope:
            if tn not in TOOL_SET:
                errors.append(tt("v.tools_scope", tn))
        scope_set = set(scope)
        used_set = set(used_tools)
        # Все вызванные должны быть в scope
        for tn in sorted(used_set - scope_set):
            errors.append(tt("v.tool_called_not_in_scope", tn))
        # Все объявленные в scope должны хотя бы раз использоваться
        for tn in sorted(scope_set - used_set):
            if tn in TOOL_SET:  # известные инструменты
                errors.append(tt("v.tool_in_scope_unused", tn))

    return errors


def _normalize_text(s: str) -> str:
    """Лёгкая нормализация для сравнения 'тривиальная одинаковость'."""
    return " ".join(s.lower().split())


# Казахские «маркеры»: спец-буквы, не встречающиеся в русском
_KAZAKH_LETTERS = set("әғіңөүұқһӘҒІҢӨҮҰҚҺ")

# Русские слова-маркеры, которые не должны встречаться в казахском тексте
_RUSSIAN_MARKERS = re.compile(
    r"\b(что|это|как|где|почему|если|тогда|пожалуйста|спасибо|здравствуйте|хочу|могу)\b",
    re.IGNORECASE,
)


def validate_warnings(form_data: dict, turns: list, archetype: str) -> list[str]:
    """
    Возвращает список ПРЕДУПРЕЖДЕНИЙ (не блокируют сохранение).
    Эти проверки полезны, но не строго требуются инструкцией v3.
    """
    warnings: list[str] = []

    # П.8: thinking_kk рекомендуется для бизнес-сценария
    if archetype == "business_scenario":
        for i, tu in enumerate(turns, 1):
            if tu["role"] == "ASSISTANT" and not tu.get("thinking_kk", "").strip():
                warnings.append(tt("w.business_thinking_recommended", i))

    # П.9: проверка казахскости текста в _kk полях
    for i, tu in enumerate(turns, 1):
        for field, label in (("content_kk", "Content"), ("thinking_kk", "Thinking")):
            text = (tu.get(field) or "").strip()
            if not text:
                continue
            # Игнорируем короткие реплики (< 30 символов) и поля TOOL_RESULT (там JSON)
            if tu.get("role") == "TOOL_RESULT":
                continue
            if len(text) < 30:
                continue
            has_kazakh = any(ch in _KAZAKH_LETTERS for ch in text)
            has_russian = bool(_RUSSIAN_MARKERS.search(text))
            if has_russian and not has_kazakh:
                warnings.append(tt("w.text_looks_russian", i, label))

    return warnings


def check_duplicate_unit_code(user_id: int, unit_code: str, exclude_id: int | None = None) -> bool:
    """П.10: проверяет, нет ли уже единицы с таким же unit_code у этого аннотатора."""
    db = get_db()
    if exclude_id is not None:
        row = db.execute(
            "SELECT 1 FROM units WHERE user_id = ? AND unit_code = ? AND id != ? LIMIT 1",
            (user_id, unit_code, exclude_id),
        ).fetchone()
    else:
        row = db.execute(
            "SELECT 1 FROM units WHERE user_id = ? AND unit_code = ? LIMIT 1",
            (user_id, unit_code),
        ).fetchone()
    return row is not None


# ─────────────────────────────────────────────────────────────────────────────
# Экспорт DOCX по шаблону v3
# ─────────────────────────────────────────────────────────────────────────────

def render_unit_text(unit: dict) -> list[str]:
    """Возвращает список строк (параграфов) одной единицы по шаблону 3.1."""
    lines: list[str] = []
    L = lines.append

    L("============================")
    L(f"UNIT {unit['unit_code']}")
    L("============================")
    L("")
    L(f"[ARCHETYPE]        {unit['archetype']}")
    if unit.get("domain"):
        L(f"[DOMAIN]           {unit['domain']}")
    if unit.get("cultural_anchor"):
        L(f"[CULTURAL_ANCHOR]  {unit['cultural_anchor']}")
    if unit.get("business_use_case"):
        L(f"[BUSINESS_USE_CASE] {unit['business_use_case']}")
    if unit.get("complexity"):
        L(f"[COMPLEXITY]       {unit['complexity']}")
    if unit.get("tools_in_scope"):
        L(f"[TOOLS_IN_SCOPE]   {unit['tools_in_scope']}")
    L("")

    if unit.get("scenario_kk"):
        L("--- СЦЕНАРИЙ ---")
        L(unit["scenario_kk"])
        L("")

    turns = json.loads(unit["turns_json"])
    for i, t in enumerate(turns, 1):
        role = t["role"]
        L(f"--- TURN {i} [{role}] ---")
        if role == "ASSISTANT" and t.get("thinking_kk", "").strip():
            L("Thinking (Kazakh):")
            L(t["thinking_kk"].rstrip())
        if t.get("content_kk", "").strip():
            label = "Content:" if role == "TOOL_RESULT" else "Content (Kazakh):"
            L(label)
            L(t["content_kk"].rstrip())
        if role == "ASSISTANT" and t.get("tool_name", "").strip():
            L("Tool call:")
            L(f"  name: {t['tool_name']}")
            if t.get("tool_args", "").strip():
                L(f"  args: {t['tool_args'].strip()}")
            if t.get("tool_mock_result", "").strip():
                L(f"  expected mock_result: {t['tool_mock_result'].strip()}")
        L("")

    if unit["archetype"] == "preference" and unit.get("preference_json"):
        pref = json.loads(unit["preference_json"])
        L("=== PREFERENCE ===")
        L(f"Good turn idx: {pref.get('good_turn_idx', '')}")
        L(f"Bad alternative (Kazakh): {pref.get('bad_alternative_kk', '').rstrip()}")
        L(f"Rejection reason: {pref.get('bad_reason', '')}")
        L("")

    L("=== METADATA ===")
    L(f"Source: {unit['source']}")
    L(f"License: {unit['license']}")
    L(f"Annotator: {unit['annotator_name']}")
    if unit.get("reviewer"):
        L(f"Reviewer: {unit['reviewer']}")
    L(f"Started at: {unit['started_at']}")
    L(f"Completed at: {unit['completed_at']}")
    L("")

    return lines


def build_docx(units: list[dict], author: str) -> BytesIO:
    """Собирает DOCX-файл с одной или несколькими единицами.
    Свойство Author документа = имя аннотатора (требование п. 3.2)."""
    doc = Document()

    # Устанавливаем core properties
    doc.core_properties.author = author
    doc.core_properties.last_modified_by = author
    doc.core_properties.title = "KazLLM-2026H2 batch"
    doc.core_properties.created = dt.datetime.utcnow()

    for idx, unit in enumerate(units):
        if idx > 0:
            doc.add_paragraph("")
            doc.add_paragraph("============================")
            doc.add_paragraph("============================")
            doc.add_paragraph("")
        for line in render_unit_text(unit):
            doc.add_paragraph(line)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


# ─────────────────────────────────────────────────────────────────────────────
# Маршруты: авторизация
# ─────────────────────────────────────────────────────────────────────────────

@app.route("/register", methods=["GET", "POST"])
def register():
    if current_user():
        return redirect(url_for("index"))
    if request.method == "POST":
        username = request.form.get("username", "").strip().lower()
        full_name = request.form.get("full_name", "").strip()
        password = request.form.get("password", "")
        password2 = request.form.get("password2", "")

        errors = []
        if not USERNAME_RE.match(username):
            errors.append(tt("reg.err_username"))
        if not full_name or len(full_name) < 2:
            errors.append(tt("reg.err_full_name"))
        if len(password) < 8:
            errors.append(tt("reg.err_password_short"))
        if password != password2:
            errors.append(tt("reg.err_password_mismatch"))

        db = get_db()
        if not errors and db.execute(
            "SELECT 1 FROM users WHERE username = ?", (username,)
        ).fetchone():
            errors.append(tt("reg.err_taken"))

        if errors:
            for e in errors:
                flash(e, "error")
            return render_template(
                "register.html", username=username, full_name=full_name
            )

        db.execute(
            "INSERT INTO users(username, full_name, password_hash) VALUES (?,?,?)",
            (username, full_name, generate_password_hash(password)),
        )
        db.commit()
        flash(tt("reg.created"), "ok")
        return redirect(url_for("login"))

    return render_template("register.html")


@app.route("/login", methods=["GET", "POST"])
def login():
    if current_user():
        return redirect(url_for("index"))
    if request.method == "POST":
        username = request.form.get("username", "").strip().lower()
        password = request.form.get("password", "")
        db = get_db()
        row = db.execute(
            "SELECT * FROM users WHERE username = ?", (username,)
        ).fetchone()
        if row and check_password_hash(row["password_hash"], password):
            session.clear()
            session["user_id"] = row["id"]
            return redirect(url_for("index"))
        flash(tt("login.bad"), "error")
    return render_template("login.html")


@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))


# ─────────────────────────────────────────────────────────────────────────────
# Маршруты: основное приложение
# ─────────────────────────────────────────────────────────────────────────────

@app.route("/")
@login_required
def index():
    db = get_db()
    units = db.execute(
        "SELECT id, unit_code, archetype, domain, created_at "
        "FROM units WHERE user_id = ? ORDER BY id DESC LIMIT 200",
        (session["user_id"],),
    ).fetchall()
    return render_template("index.html", units=units)


@app.route("/help")
@login_required
def help_page():
    return render_template("help.html")


@app.route("/examples")
@login_required
def examples_page():
    """Страница со всеми примерами — по ключам в EXAMPLES."""
    lang = get_lang()
    items = []
    for key, ex_full in EXAMPLES.items():
        items.append({
            "key": key,
            "archetype": ex_full["data"]["archetype"],
            "title": ex_full.get(f"title_{lang}", ex_full.get("title_ru", key)),
            "explanation": ex_full.get(f"explanation_{lang}", ex_full.get("explanation_ru", "")),
            "data": ex_full["data"],
        })
    # Подгружаем правила архетипов для отображения над примерами
    from archetype_rules import all_rules
    rules = all_rules(lang)
    return render_template("examples.html", items=items, rules=rules)


@app.route("/unit/new/from-example/<example_key>")
@login_required
def unit_new_from_example(example_key):
    """Открывает форму создания, предзаполненную данными из примера.
    example_key — ключ примера в EXAMPLES (не обязательно совпадает с archetype,
    т.к. для одного архетипа может быть несколько примеров).
    Unit_code предлагается со следующим свободным номером для этого пользователя."""
    ex = EXAMPLES.get(example_key)
    if not ex:
        flash(tt("examples.unknown"), "error")
        return redirect(url_for("examples_page"))

    d = ex["data"]
    # Префикс примера: ly-ds-2026h2-business-001 → ly-ds-2026h2-business
    prefix = d["unit_code"].rsplit("-", 1)[0]
    db = get_db()
    rows = db.execute(
        "SELECT unit_code FROM units WHERE user_id = ? AND unit_code LIKE ?",
        (session["user_id"], prefix + "-%"),
    ).fetchall()
    used_numbers = set()
    for row in rows:
        m = re.match(rf"^{re.escape(prefix)}-(\d+)$", row["unit_code"])
        if m:
            used_numbers.add(int(m.group(1)))
    next_num = 1
    while next_num in used_numbers or next_num == 1:
        next_num += 1
        if next_num > 999999:
            break
    suggested_code = f"{prefix}-{next_num:03d}"

    unit_stub = {
        "id": None,
        "unit_code":         suggested_code,
        "archetype":         d["archetype"],
        "domain":            d["domain"] or "",
        "cultural_anchor":   d["cultural_anchor"] or "",
        "business_use_case": d["business_use_case"] or "",
        "complexity":        d["complexity"] or "",
        "tools_in_scope":    d["tools_in_scope"] or "",
        "scenario_kk":       d["scenario_kk"] or "",
        "source":            d["source"],
        "license":           d["license"],
        "reviewer":          d["reviewer"] or "",
        "started_at":        d["started_at"],
        "completed_at":      d["completed_at"],
    }

    return render_template(
        "unit_form.html",
        unit=unit_stub,
        turns=d["turns"],
        preference=d.get("preference") or {},
        archetypes=_localized_archetypes(),
        domains=DOMAINS,
        anchors=CULTURAL_ANCHORS,
        complexities=COMPLEXITIES,
        licenses=LICENSES,
        rejection_reasons=_localized_reasons(),
        tool_library=TOOL_LIBRARY,
        suggested_started=d["started_at"],
        loaded_from_example=example_key,
    )


@app.route("/unit/new", methods=["GET", "POST"])
@login_required
def unit_new():
    if request.method == "POST":
        return _save_unit(unit_id=None)
    # Стартовое время по умолчанию — сейчас, по Алматы (+06:00)
    now = dt.datetime.now(dt.timezone(dt.timedelta(hours=6)))
    return render_template(
        "unit_form.html",
        unit=None,
        turns=[{"role": "USER", "content_kk": ""}],
        preference={},
        archetypes=_localized_archetypes(),
        domains=DOMAINS,
        anchors=CULTURAL_ANCHORS,
        complexities=COMPLEXITIES,
        licenses=LICENSES,
        rejection_reasons=_localized_reasons(),
        tool_library=TOOL_LIBRARY,
        suggested_started=now.strftime("%Y-%m-%dT%H:%M:%S+06:00"),
    )


@app.route("/unit/<int:unit_id>/edit", methods=["GET", "POST"])
@login_required
def unit_edit(unit_id):
    db = get_db()
    row = db.execute(
        "SELECT * FROM units WHERE id = ? AND user_id = ?",
        (unit_id, session["user_id"]),
    ).fetchone()
    if not row:
        abort(404)

    if request.method == "POST":
        return _save_unit(unit_id=unit_id)

    turns = json.loads(row["turns_json"])
    preference = json.loads(row["preference_json"]) if row["preference_json"] else {}
    return render_template(
        "unit_form.html",
        unit=dict(row),
        turns=turns,
        preference=preference,
        archetypes=_localized_archetypes(),
        domains=DOMAINS,
        anchors=CULTURAL_ANCHORS,
        complexities=COMPLEXITIES,
        licenses=LICENSES,
        rejection_reasons=_localized_reasons(),
        tool_library=TOOL_LIBRARY,
        suggested_started=row["started_at"],
    )


@app.route("/unit/<int:unit_id>/delete", methods=["POST"])
@login_required
def unit_delete(unit_id):
    db = get_db()
    db.execute(
        "DELETE FROM units WHERE id = ? AND user_id = ?",
        (unit_id, session["user_id"]),
    )
    db.commit()
    flash(tt("index.deleted"), "ok")
    return redirect(url_for("index"))


def _collect_turns_from_form(form) -> list[dict]:
    """Собирает туры из плоской формы (turn_<i>_<field>).

    Особенность: в форме два разных поля для содержимого тура:
      - turn_N_content_kk_text  — казахский текст для USER/ASSISTANT
      - turn_N_content_kk_json  — JSON-результат для TOOL_RESULT
    На сервере мы объединяем их в одно поле content_kk в зависимости от роли.
    """
    turns: list[dict] = []
    i = 0
    while True:
        prefix = f"turn_{i}_"
        if f"{prefix}role" not in form:
            break
        role = form.get(f"{prefix}role", "").strip()
        if role == "TOOL_RESULT":
            content = form.get(f"{prefix}content_kk_json", "")
        else:
            content = form.get(f"{prefix}content_kk_text", "")
        t = {
            "role":             role,
            "content_kk":       content,
            "thinking_kk":      form.get(f"{prefix}thinking_kk", ""),
            "tool_name":        form.get(f"{prefix}tool_name", "").strip() if role == "ASSISTANT" else "",
            "tool_args":        form.get(f"{prefix}tool_args", "") if role == "ASSISTANT" else "",
            "tool_mock_result": form.get(f"{prefix}tool_mock_result", "") if role == "ASSISTANT" else "",
        }
        turns.append(t)
        i += 1
    return turns


def _save_unit(unit_id):
    form = request.form
    turns = _collect_turns_from_form(form)

    archetype = form.get("archetype", "")
    preference = None
    if archetype == "preference":
        preference = {
            "good_turn_idx":      form.get("pref_good_turn_idx", ""),
            "bad_alternative_kk": form.get("pref_bad_alternative_kk", ""),
            "bad_reason":         form.get("pref_bad_reason", ""),
        }

    errors = validate_unit(form, turns, preference)

    # П.10: проверка дубликата unit_code у этого пользователя
    unit_code = form.get("unit_code", "").strip()
    if unit_code and UNIT_CODE_RE.match(unit_code):
        if check_duplicate_unit_code(session["user_id"], unit_code, exclude_id=unit_id):
            errors.append(tt("v.duplicate_unit_code", unit_code))

    # Предупреждения (не блокируют, показываем пользователю отдельно)
    warnings = validate_warnings(form, turns, archetype)

    if errors:
        for e in errors:
            flash(e, "error")
        for w in warnings:
            flash(w, "warning")
        # Перерисовываем форму с введёнными данными
        unit_stub = {
            "id": unit_id,
            "unit_code":         form.get("unit_code", ""),
            "archetype":         archetype,
            "domain":            form.get("domain", ""),
            "cultural_anchor":   form.get("cultural_anchor", ""),
            "business_use_case": form.get("business_use_case", ""),
            "complexity":        form.get("complexity", ""),
            "tools_in_scope":    form.get("tools_in_scope", ""),
            "scenario_kk":       form.get("scenario_kk", ""),
            "source":             form.get("source", ""),
            "license":            form.get("license", ""),
            "reviewer":           form.get("reviewer", ""),
            "started_at":         form.get("started_at", ""),
            "completed_at":       form.get("completed_at", ""),
        }
        return render_template(
            "unit_form.html",
            unit=unit_stub,
            turns=turns,
            preference=preference or {},
            archetypes=_localized_archetypes(),
            domains=DOMAINS,
            anchors=CULTURAL_ANCHORS,
            complexities=COMPLEXITIES,
            licenses=LICENSES,
            rejection_reasons=_localized_reasons(),
            tool_library=TOOL_LIBRARY,
            suggested_started=form.get("started_at", ""),
        )

    db = get_db()
    fields = (
        form.get("unit_code", "").strip(),
        archetype,
        form.get("domain", "").strip() or None,
        form.get("cultural_anchor", "").strip() or None,
        form.get("business_use_case", "").strip() or None,
        form.get("complexity", "").strip() or None,
        form.get("tools_in_scope", "").strip() or None,
        form.get("scenario_kk", "").strip() or None,
        json.dumps(turns, ensure_ascii=False),
        json.dumps(preference, ensure_ascii=False) if preference else None,
        form.get("source", "").strip(),
        form.get("license", "").strip(),
        form.get("reviewer", "").strip() or None,
        form.get("started_at", "").strip(),
        form.get("completed_at", "").strip(),
    )

    if unit_id is None:
        db.execute(
            """INSERT INTO units(
                user_id, unit_code, archetype, domain, cultural_anchor,
                business_use_case, complexity, tools_in_scope, scenario_kk,
                turns_json, preference_json, source, license, reviewer,
                started_at, completed_at
            ) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
            (session["user_id"], *fields),
        )
    else:
        db.execute(
            """UPDATE units SET
                unit_code=?, archetype=?, domain=?, cultural_anchor=?,
                business_use_case=?, complexity=?, tools_in_scope=?,
                scenario_kk=?, turns_json=?, preference_json=?,
                source=?, license=?, reviewer=?, started_at=?, completed_at=?
               WHERE id=? AND user_id=?""",
            (*fields, unit_id, session["user_id"]),
        )
    db.commit()
    flash(tt("form.saved_ok"), "ok")
    for w in warnings:
        flash(w, "warning")
    return redirect(url_for("index"))


# ─────────────────────────────────────────────────────────────────────────────
# Экспорт
# ─────────────────────────────────────────────────────────────────────────────

@app.route("/export", methods=["GET", "POST"])
@login_required
def export():
    db = get_db()
    user = current_user()
    if request.method == "POST":
        ids = request.form.getlist("unit_ids")
        if not ids:
            flash(tt("export.must_select"), "error")
            return redirect(url_for("export"))

        # Только свои единицы
        placeholders = ",".join("?" * len(ids))
        rows = db.execute(
            f"SELECT * FROM units WHERE user_id = ? AND id IN ({placeholders})",
            (session["user_id"], *ids),
        ).fetchall()

        units = []
        for r in rows:
            d = dict(r)
            d["annotator_name"] = user["full_name"]
            units.append(d)

        bio = build_docx(units, author=user["full_name"])
        ym = dt.datetime.now().strftime("%Y_%m")
        # Категория — берём по первой единице, если все одинаковые
        cats = {u["archetype"] for u in units}
        cat = next(iter(cats)) if len(cats) == 1 else "mixed"
        filename = f"batch_{ym}_{cat}_01.docx"

        return send_file(
            bio,
            as_attachment=True,
            download_name=filename,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    units = db.execute(
        "SELECT id, unit_code, archetype, domain, created_at "
        "FROM units WHERE user_id = ? ORDER BY id DESC",
        (session["user_id"],),
    ).fetchall()
    return render_template("export.html", units=units)


# ─────────────────────────────────────────────────────────────────────────────
# JSON API для проверки на лету (используется JS-валидатором)
# ─────────────────────────────────────────────────────────────────────────────

@app.route("/api/validate-iin", methods=["POST"])
@login_required
def api_validate_iin():
    iin = (request.json or {}).get("iin", "")
    return jsonify({"valid": validate_iin(iin)})


@app.route("/api/tool-skeleton/<tool_name>")
@login_required
def api_tool_skeleton(tool_name):
    """Возвращает JSON-скелеты input и output для указанного инструмента.
    Используется JS-кодом формы, чтобы подставить шаблон при выборе инструмента."""
    schema = TOOL_SCHEMAS.get(tool_name)
    if not schema:
        return jsonify({"error": "unknown tool"}), 404
    spec = get_input_spec(tool_name) or {}
    return jsonify({
        "tool": tool_name,
        "domain": schema.get("domain", []),
        "input_skeleton":  empty_input_for(schema),
        "output_skeleton": empty_output_for(schema),
        "input_required":  spec.get("required", []),
        "input_optional":  spec.get("optional", []),
    })


@app.route("/api/tool-form/<tool_name>")
@login_required
def api_tool_form(tool_name):
    """Возвращает описание UI-формы для инструмента: список полей с лейблами,
    типами и hint-ами на текущем языке. JS использует это для рендера полей
    вместо того, чтобы просить аннотатора писать JSON руками."""
    schema = TOOL_SCHEMAS.get(tool_name)
    if not schema:
        return jsonify({"error": "unknown tool"}), 404
    return jsonify(schema_for_form(schema, lang=get_lang()))


# ─────────────────────────────────────────────────────────────────────────────
# Точка входа
# ─────────────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    init_db()
    print("=" * 60)
    print("KazLLM Annotator — локальный сервер")
    print("Открой в браузере: http://127.0.0.1:5000")
    print("Чтобы остановить — Ctrl+C")
    print("=" * 60)
    app.run(host="127.0.0.1", port=5000, debug=False)
